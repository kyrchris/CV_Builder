import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def remove_space_after_paragraph(paragraph):
    paragraph.style.paragraph_format.space_after = Pt(0)


def set_narrow_margins(document):
    # Access the first section of the document
    section = document.sections[0]

    # Set narrow margins (e.g., 1.27 cm for all sides)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)


# Create Word Document
document = Document()
set_narrow_margins(document)

# Set default font and text size
style = document.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

#####-----SECTION 1----#####

##Personal info section
par1 = document.add_paragraph()
while True:
    name_input = input("Please enter your full name: ")

    if not name_input.isdigit():  # Check if the input is not a number
        break
    else:
        print("Please enter a valid string. Numbers are not allowed.")

# Add the user's input to the run as a string
name = par1.add_run(str(name_input))

name.font.size = Pt(16)
name.bold = True
par1.add_run().add_break()  # changes line on the document

while True:
    Country = input("Please enter your country: ")

    if not Country.isdigit():  # Check if the input is not a number
        break
    else:
        print("Please enter a valid string. Numbers are not allowed.")

while True:
    City = input("Please enter your city: ")

    if not City.isdigit():  # Check if the input is not a number
        break
    else:
        print("Please enter a valid string. Numbers are not allowed.")

Phone = input("Please enter your phone number (including the country code): ")
line2 = par1.add_run(City + ", " + Country + " " + Phone + " , ")
contact = input("Please enter your e-mail: ")
par1.add_run(contact)
line2.font.size = Pt(12)
par1.add_run().add_break()
online_contact = input("Do you want to add an online profile? (e.g. Portfolio website or Linked in profile)? Yes/No: ")
while online_contact.lower() not in ["yes", "no"]:
    online_contact = input(
        "Do you want to add an online profile? (e.g. Portfolio website or Linked in profile)? Yes/No: ")

if online_contact.lower() == "yes":
    contact_type = input("Is it a Portfolio website or a LinkdeIn Profile?: ")
    while contact_type.lower() not in ["linkedin", "portfolio"]:
        contact_type = input("Is it a portfolio or a LinkdeIn Profile?: ")
    if contact_type.lower() == "portfolio":
        par1.add_run("Portfolio: ")
        portfolio = input("Please enter your Portfolio's url:")
        par1.add_run(portfolio)
    elif contact_type.lower() == "linkedin":
        par1.add_run("LinkedIn: ")
        portfolio = input("Please enter your LinkedIn's url:")
        par1.add_run(portfolio)
    else:
        contact_type = input("What kind of online profile do you want to add?: ")
        par1.add_run(contact_type + ": ")
        url = input("Please enter a url:")
        par1.add_run(url)

par1.alignment = WD_ALIGN_PARAGRAPH.CENTER

#This line seperates each section
document.add_paragraph(
    "__________________________________________________________________________________________________")

#####---SECTION 2----#####

#Work experience section
par2 = document.add_paragraph("")
line1 = par2.add_run("Work experience")
line1.font.size = Pt(14)
line1.bold = True
par2.add_run().add_break()

#Loop to ensure that the user adds a number of jobs that is integer and positive
while True:
    try:
        num_Jobs = int(input("How many Jobs do you want to add?: "))
        if num_Jobs >= 0:
            break
        else:
            print("Please enter a positive value larger than zero.")
    except ValueError:
        print("Please enter a numeric value.")

#accounting for no work experience
if num_Jobs == 0:
    par2 = document.add_paragraph("No work experience")

#accounting for one job only
elif num_Jobs == 1:
    par2 = document.add_paragraph(None)
    Job1 = input("Please enter your current job title and Company in brackets: ")
    date_start1 = input("When did you start at this job?: ")
    end_date1 = "Present"
    line2 = par2.add_run(Job1 + " (" + date_start1 + "-" + end_date1 + ")")
    line2.bold = True
    line2.font.size = Pt(12)
    while True:
        try:
            num_bulletpoints = int(
                input("How many bullet-points do you want to add to describe what you did as a " + Job1 + "? "))
            break
        except ValueError:
            print("Please enter a numeric value.")

    list = []

    for _ in range(num_bulletpoints):
        bullet_point = input("Enter a bullet point: ")
        list.append(bullet_point)
    for l in list:
        par = document.add_paragraph(None, style='List Bullet')
        line = par.add_run(l)
        line.font.size = Pt(11)

    document.add_paragraph(None)
else:
    par2 = document.add_paragraph(None)
    Job1 = input("Please enter your current job title and Company in brackets: ")
    date_start1 = input("When did you start at this job?: ")
    end_date1 = "Present"
    line2 = par2.add_run(Job1 + " (" + date_start1 + "-" + end_date1 + ")")
    line2.bold = True
    line2.font.size = Pt(12)
    while True:
        try:
            num_bulletpoints = int(
                input("How many bullet-points do you want to add to describe what you did as a " + Job1 + "? "))
            break
        except ValueError:
            print("Please enter a numeric value.")

    list = []

    for _ in range(num_bulletpoints):
        bullet_point = input("Enter a bullet point: ")
        list.append(bullet_point)
    for l in list:
        par = document.add_paragraph(None, style='List Bullet')
        line = par.add_run(l)
        line.font.size = Pt(11)

    document.add_paragraph(None)
    num_Jobs = num_Jobs - 1
    for _ in range(num_Jobs):
        par2 = document.add_paragraph(None)
        Job1 = input("Please enter your previous job: ")
        date_start1 = input("When did you start at this job?: ")
        end_date1 = input("Please enter the date you left this job: ")

        line2 = par2.add_run(Job1 + " (" + date_start1 + "-" + end_date1 + ")")
        line2.bold = True
        line2.font.size = Pt(12)
        while True:
            try:
                num_bulletpoints = int(
                    input("How many bullet-points do you want to add to describe what you did as a " + Job1 + "? "))
                break
            except ValueError:
                print("Please enter a numeric value.")

        list = []

        for _ in range(num_bulletpoints):
            bullet_point = input("Enter a bullet point: ")
            list.append(bullet_point)
        for l in list:
            par = document.add_paragraph(None, style='List Bullet')
            line = par.add_run(l)
            line.font.size = Pt(11)

        document.add_paragraph(None)

document.add_paragraph(
    "__________________________________________________________________________________________________")

####---SECTION 3---####

#Education section
par2 = document.add_paragraph("")
line1 = par2.add_run("Education")
line1.font.size = Pt(14)
line1.bold = True
par2.add_run().add_break()
bachelors = input("Do you have a bachelor's degree? yes/no: ")

while bachelors.lower() not in ["yes", "no"]:
    bachelors = input("Do you have a bachelor's degree? yes/no: ")

if bachelors == "yes":
    uni = input("Where did you study?")
    bachelors_degree = input("Please enter the name of your bachelors degree: ")
    bachelors_start = input("When did you start your bachelors degree?: ")
    bachelors_end = input("When did you finish your bachelors degree?: ")

    line2 = par2.add_run("BSc, " + bachelors_degree + ", " + uni + " (" + bachelors_start + "-" + bachelors_end + ")")
    line2.bold = True
    line2.font.size = Pt(12)
    par2.add_run().add_break()
    grade = input("What was your grade?: ")
    line3 = par2.add_run("Grade: ")
    line3.bold = True
    line3.font.size = Pt(11)
    line4 = par2.add_run(grade)
    line4.font.size = Pt(11)
    par2.add_run().add_break()

par2 = document.add_paragraph("")

masters = input("Do you have a master's degree? yes/no: ")

while masters.lower() not in ["yes", "no"]:
    masters = input("Do you have a bachelor's degree? yes/no: ")

if masters == "yes":
    uni = input("Where did you study?")
    masters_degree = input("Please enter the name of your masters degree: ")
    masters_start = input("When did you start your masters degree?: ")

    masters_end = input("When did you finish your masters degree?: ")

    line2 = par2.add_run("MSc, " + masters_degree + ", " + uni + " (" + masters_start + "-" + masters_end + ")")
    line2.bold = True
    line2.font.size = Pt(12)
    par2.add_run().add_break()
    grade = input("What was your grade?: ")
    line3 = par2.add_run("Grade: ")
    line3.bold = True
    line3.font.size = Pt(11)
    line4 = par2.add_run(grade)
    line4.font.size = Pt(11)
    par2.add_run().add_break()

document.add_paragraph(
    "__________________________________________________________________________________________________")

####---SECTION 4---####

#Computer skills section
par2 = document.add_paragraph("")
remove_space_after_paragraph(par2)
line1 = par2.add_run("Computers literacy")
line1.font.size = Pt(14)
line1.bold = True
par2.add_run().add_break()

while True:
    try:
        num_Computer_Skills = int(input("How many Computer skills do you want to add? : "))
        break
    except ValueError:
        print("Please enter a numeric value.")

if num_Computer_Skills == 0:
    table = None
#This will create a 2x2 matrix including your computers literacy
elif num_Computer_Skills == 1:
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = input("Please enter the name of the software you want to add: ")
    hdr_cells[1].text = input(
        "Please enter the level you are familiar with the software (Expert/Experienced/Skillful): ")
else:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False  # Disable autofit to set column widths manually

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Software"
    hdr_cells[1].text = "Level"

    for _ in range(num_Computer_Skills):
        software = input("Please enter the name of the software you want to add: ")
        level = input("Please enter the level you are familiar with the software (Expert/Experienced/Skillful): ")
        row_cells = table.add_row().cells
        row_cells[0].text = software
        row_cells[1].text = level

document.add_paragraph(
    "__________________________________________________________________________________________________")

####---SECTION 5---####

#Internship Section
internships = input("Do you want to add internships? yes/no: ")
while internships.lower() not in ["yes", "no"]:
    input("Do you want to add internships? yes/no: ")
if internships == "yes":
    par2 = document.add_paragraph("")
    line1 = par2.add_run("Internships")
    line1.font.size = Pt(14)
    line1.bold = True
    while True:
        try:
            num_Internships = int(input("How many Internships do you want to add?: "))
            break
        except ValueError:
            print("Please enter a numeric value.")

    for _ in range(num_Internships):
        par2 = document.add_paragraph(None)
        Intern1 = input("Please enter the name of the company that you Internshiped for: ")
        date_start1 = input("When did you start at this Internship?: ")
        end_date1 = input("Please enter the date you finished your Internship: ")

        line2 = par2.add_run(Intern1 + " (" + date_start1 + "-" + end_date1 + ")")
        line2.bold = True
        line2.font.size = Pt(12)
        document.add_paragraph(None)

    document.add_paragraph(
        "__________________________________________________________________________________________________")

####---SECTION 6---####

#Languages section
par2 = document.add_paragraph("")
line1 = par2.add_run("Languages")
line1.font.size = Pt(14)
line1.bold = True
par2 = document.add_paragraph("")
while True:
    try:
        num_Languages = int(input("How many languages do you speak? : "))
        break
    except ValueError:
        print("Please enter a numeric value.")
while num_Languages == 0:
    try:
        num_Languages = int(input("How many languages do you speak? : "))
        if num_Languages <= 0:
            print("Please enter a positive value larger than 0.")
            num_Languages = 0  # Reset num_Languages to trigger the loop again
    except ValueError:
        print("Please enter a valid positive integer.")

while num_Languages < 0:
    try:
        num_Languages = int(input("How many languages do you speak? : "))
        if num_Languages <= 0:
            print("Please enter a positive value larger than 0.")
            num_Languages = 0  # Reset num_Languages to trigger the loop again
    except ValueError:
        print("Please enter a valid positive integer.")

if num_Languages == 1:
    language = input("What is your native language?: ")
    lin2 = par2.add_run(language + " - Native")
else:
    native_language = input("What is your native language?: ")
    lin2 = par2.add_run(native_language + " - Native          ")
    num_Languages = num_Languages - 1
    for _ in range(num_Languages):
        language = input("What other language do you speak?: ")
        level_of_knoweledge = input("At what level do you speak this language?: ")
        line2 = par2.add_run(language + " - " + level_of_knoweledge + "          ")

    line2.font.size = Pt(11)

#Where to save the CV that you just created.
folder_destination = "add the destination path that you want to save your CV" + str(name_input) + ".docx"
document.save(folder_destination)